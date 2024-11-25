import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
import io

def load_data():
    return pd.read_csv('Testing Metrics.csv')

def extract_age_group(age_str):
    return int(age_str.replace('u', ''))

def calculate_percentile(value, series):
    return round(100 * (series <= value).mean())

def calculate_age_based_stats(player_data, df):
    # Check for NaN values
    required_fields = ['age'] + [f'{metric} {i}' for metric in ['Bat Speed', 'Rot. Acc.', 'Exit Velo', 'VBA'] for i in range(1,6)]
    
    valid_bat_speeds = [float(player_data[f'Bat Speed {i}']) for i in range(1,6) if not pd.isna(player_data[f'Bat Speed {i}'])]
    valid_rot_accs = [float(player_data[f'Rot. Acc. {i}']) for i in range(1,6) if not pd.isna(player_data[f'Rot. Acc. {i}'])]
    valid_exit_velos = [float(player_data[f'Exit Velo {i}']) for i in range(1,6) if not pd.isna(player_data[f'Exit Velo {i}'])]
    valid_vbas = [float(player_data[f'VBA {i}']) for i in range(1,6) if not pd.isna(player_data[f'VBA {i}'])]
    
    if not valid_bat_speeds or not valid_rot_accs or not valid_exit_velos or not valid_vbas:
        st.error("Not enough valid swing data to calculate statistics.")
        return None
    
    try:
        df['age_group'] = df['age'].apply(extract_age_group)
        player_age_group = extract_age_group(player_data['age'])
        age_group_df = df[df['age_group'] == player_age_group]
        
        age_bat_speeds = age_group_df[[f'Bat Speed {i}' for i in range(1,6)]].mean(axis=1)
        age_rot_accs = age_group_df[[f'Rot. Acc. {i}' for i in range(1,6)]].mean(axis=1)
        age_exit_velos = age_group_df[[f'Exit Velo {i}' for i in range(1,6)]].mean(axis=1)
        
        age_max_bat_speeds = age_group_df[[f'Bat Speed {i}' for i in range(1,6)]].max(axis=1)
        age_max_rot_accs = age_group_df[[f'Rot. Acc. {i}' for i in range(1,6)]].max(axis=1)
        age_max_exit_velos = age_group_df[[f'Exit Velo {i}' for i in range(1,6)]].max(axis=1)

        vba_high = sum(1 for x in valid_vbas if x > -24)
        vba_low = sum(1 for x in valid_vbas if x < -45)
        avg_rot_acc = np.mean(valid_rot_accs)
        avg_vba = np.mean(valid_vbas)
        
        vba_issue = vba_high >= 3 or vba_low >= 3
        rot_issue = avg_rot_acc < 7.0
        decel_issue = not (vba_issue or rot_issue)
        
        return {
            'age_group': player_age_group,
            'bat_speed': {
                'avg': np.mean(valid_bat_speeds),
                'percentile': calculate_percentile(np.mean(valid_bat_speeds), age_bat_speeds)
            },
            'max_bat_speed': {
                'value': np.max(valid_bat_speeds),
                'percentile': calculate_percentile(np.max(valid_bat_speeds), age_max_bat_speeds)
            },
            'rot_acc': {
                'avg': avg_rot_acc,
                'percentile': calculate_percentile(np.mean(valid_rot_accs), age_rot_accs)
            },
            'max_rot_acc': {
                'value': np.max(valid_rot_accs),
                'percentile': calculate_percentile(np.max(valid_rot_accs), age_max_rot_accs)
            },
            'exit_velo': {
                'avg': np.mean(valid_exit_velos),
                'percentile': calculate_percentile(np.mean(valid_exit_velos), age_exit_velos)
            },
            'max_exit_velo': {
                'value': np.max(valid_exit_velos),
                'percentile': calculate_percentile(np.max(valid_exit_velos), age_max_exit_velos)
            },
            'swing_issues': {
                'vba_high': vba_high,
                'vba_low': vba_low,
                'vba_issue': vba_issue,
                'rot_issue': rot_issue,
                'decel_issue': decel_issue,
                'avg_vba': avg_vba
            }
        }
    except Exception as e:
        st.error(f"Error calculating stats: {str(e)}")
        return None

def generate_report(player_data, stats):
    template_path = 'templates/'
    
    if stats['swing_issues']['vba_issue']:
        doc = Document(template_path + 'VBA_template.docx')
        summary = f"{player_data['First Name']}'s exit velocity average was {stats['exit_velo']['avg']:.1f}mph and their swing test showed {stats['swing_issues']['vba_high']} swings above -24°. Their average VBA was {stats['swing_issues']['avg_vba']:.1f}°. Ideally, we want to see their bat more vertical. Once achieved, it will allow them to stay \"on plane\" with the ball longer, which enables them to hit the ball hard when their timing is off. The drills below will help, I recommend 3 sets of 8 reps of each."
    elif stats['swing_issues']['rot_issue']:
        doc = Document(template_path + 'RotAcc_template.docx')
        summary = f"{player_data['First Name']}'s exit velocity average was {stats['exit_velo']['avg']:.1f}mph. They were placed in this program because their Rotational Acceleration results averaged {stats['rot_acc']['avg']:.1f}g's (Ideally, we want this 15+). What this means is that they are rotating out of order (sequence), which will reduce their barrel accuracy & rotational speed. The drills listed below will help, I recommend 3 sets of 8 reps of each."
    else:
        doc = Document(template_path + 'Decel_template.docx')
        summary = f"{player_data['First Name']}'s exit velocity average was {stats['exit_velo']['avg']:.1f}mph. Based on the swing test results, an area they need to focus on is deceleration. In order for one body part to speed up the other needs to hit the brakes. Once achieved, their body will rotate faster and more efficiently. The drills listed below will help, I recommend 3 sets of 8-10 reps each."

    # Replace header content
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if "{{PLAYER_NAME}}" in run.text:
                    run.text = run.text.replace("{{PLAYER_NAME}}", 
                        f"{player_data['First Name']} {player_data['Last Name']}")
                if "{{SUMMARY}}" in run.text:
                    run.text = run.text.replace("{{SUMMARY}}", summary)

    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream
    
st.set_page_config(page_title="Baseball Metrics Dashboard", layout="wide")
st.title("Player Metrics Dashboard")

df = load_data()

search = st.text_input("", placeholder="Search by name...")
if search:
    matches = df[df['First Name'].str.contains(search, case=False) | df['Last Name'].str.contains(search, case=False)]
    if not matches.empty:
        player_options = [f"{player['First Name']} {player['Last Name']} ({player['age']})" for _, player in matches.iterrows()]
        selected_player = st.selectbox("Select a player", player_options)
        
        player_index = player_options.index(selected_player)
        player = matches.iloc[player_index]
        
        stats = calculate_age_based_stats(player, df)
        if stats is None:
            st.error("Unable to calculate stats due to missing data")
        else:
            st.header(selected_player)
            
            st.subheader("Average Metrics")
            cols = st.columns(3)
            cols[0].metric("Avg Bat Speed (mph)", f"{stats['bat_speed']['avg']:.1f}", f"{stats['bat_speed']['percentile']}%ile in {stats['age_group']}u")
            cols[1].metric("Avg Rot. Acc. (g)", f"{stats['rot_acc']['avg']:.1f}", f"{stats['rot_acc']['percentile']}%ile in {stats['age_group']}u") 
            cols[2].metric("Avg Exit Velo (mph)", f"{stats['exit_velo']['avg']:.1f}", f"{stats['exit_velo']['percentile']}%ile in {stats['age_group']}u")

            st.subheader("Max Metrics")
            cols = st.columns(3)
            cols[0].metric("Max Bat Speed (mph)", f"{stats['max_bat_speed']['value']:.1f}", f"{stats['max_bat_speed']['percentile']}%ile in {stats['age_group']}u")
            cols[1].metric("Max Rot. Acc. (g)", f"{stats['max_rot_acc']['value']:.1f}", f"{stats['max_rot_acc']['percentile']}%ile in {stats['age_group']}u")
            cols[2].metric("Max Exit Velo (mph)", f"{stats['max_exit_velo']['value']:.1f}", f"{stats['max_exit_velo']['percentile']}%ile in {stats['age_group']}u")
            
            st.subheader("Swing Issues")
            issues = stats['swing_issues']
            if issues['vba_issue']:
                message = [f"{issues['vba_high']} swings >-24°", f"{issues['vba_low']} swings <-45°"]
                st.warning("VBA Issue: " + " and ".join(message))
            elif issues['rot_issue']:
                st.warning("Rotational Acceleration Issue: Avg <7.0g")
            else:
                st.success("Deceleration Pattern")
            
            if st.button("Export Training Reports", type="primary"):
                docx_file = generate_report(player, stats)
                filename = f"{player['Last Name']}_Individualized_Hitting_Plan_{player['First Name']}_{player['age']}.docx"
                st.download_button("Download Plan", docx_file, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("No player found")
